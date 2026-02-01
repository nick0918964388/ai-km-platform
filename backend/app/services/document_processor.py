"""Document processing service for PDF, Word, Excel, and images."""
import base64
import io
import logging
import uuid
from typing import Optional, Callable

import fitz  # PyMuPDF - for image extraction
import pdfplumber  # For text extraction (better CJK support)
from docx import Document as DocxDocument
from openpyxl import load_workbook
import xlrd
from PIL import Image

from app.services import embedding as embed_service
from app.services import vector_store
from app.services import file_storage
from app.services.chunking import smart_chunk, chunk_document
from app.models.schemas import ProcessingStep

logger = logging.getLogger(__name__)

# Progress callback type: (step, progress_percent, message, chunk_count)
ProgressCallback = Callable[[ProcessingStep, int, str, Optional[int]], None]


def process_pdf(
    file_content: bytes,
    filename: str,
    use_smart_chunking: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> tuple[str, int]:
    """
    Process a PDF file and store chunks in vector store.

    Uses pdfplumber for text extraction (better Chinese support)
    and PyMuPDF for image extraction.

    Args:
        file_content: Raw PDF bytes
        filename: Original filename
        use_smart_chunking: Whether to use smart chunking
        progress_callback: Optional callback for progress updates

    Returns: (document_id, chunk_count)
    """
    document_id = str(uuid.uuid4())
    chunk_count = 0

    def emit_progress(step: ProcessingStep, progress: int, message: str, chunks: Optional[int] = None):
        if progress_callback:
            try:
                progress_callback(step, progress, message, chunks)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")

    images_data = []

    # Save original file to storage
    try:
        file_storage.save_file(document_id, filename, file_content)
        logger.info(f"Saved original file for document {document_id}")
    except Exception as e:
        logger.warning(f"Failed to save original file: {e}")

    # Step 1: Parsing (0-25%)
    emit_progress(ProcessingStep.PARSING, 0, "開始解析 PDF 文件...")

    # Extract all text using pdfplumber
    full_text = ""
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        total_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                full_text += text + "\n\n"
            # Progress: 0-20% for text extraction
            page_progress = int((page_num + 1) / total_pages * 20)
            emit_progress(ProcessingStep.PARSING, page_progress, f"解析頁面 {page_num + 1}/{total_pages}...")

    # Extract images using PyMuPDF (20-25%)
    emit_progress(ProcessingStep.PARSING, 20, "提取圖片中...")
    doc = fitz.open(stream=file_content, filetype="pdf")
    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # Convert to base64
                image_base64 = base64.b64encode(image_bytes).decode("utf-8")
                images_data.append({
                    "base64": image_base64,
                    "page": page_num + 1,
                })
            except Exception:
                continue
    doc.close()
    emit_progress(ProcessingStep.PARSING, 25, f"解析完成，找到 {len(images_data)} 張圖片")

    # Step 2: Chunking (25-50%)
    emit_progress(ProcessingStep.CHUNKING, 25, "文件分塊中...")

    # Use smart chunking for technical documents
    if use_smart_chunking and full_text.strip():
        chunks = smart_chunk(full_text, chunk_size=500, overlap=100)
    else:
        # Fallback to simple paragraph chunking
        chunks = [
            {"content": p.strip(), "metadata": {}}
            for p in full_text.split("\n\n")
            if p.strip() and len(p.strip()) > 30
        ]

    total_items = len(chunks) + len(images_data)
    emit_progress(ProcessingStep.CHUNKING, 50, f"分塊完成，共 {len(chunks)} 個文字區塊", len(chunks))

    # Step 3: Embedding (50-90%)
    emit_progress(ProcessingStep.EMBEDDING, 50, "向量化文字區塊中...")

    # Process text chunks with metadata
    if chunks:
        contents = [c["content"] for c in chunks]
        embeddings = embed_service.embed_texts(contents)
        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            vector_store.add_text_chunk(
                document_id=document_id,
                document_name=filename,
                content=chunk["content"],
                embedding=emb,
                chunk_index=idx,
                metadata=chunk.get("metadata", {}),
            )
            chunk_count += 1
            # Progress: 50-75% for text embedding
            if total_items > 0:
                embed_progress = 50 + int((idx + 1) / total_items * 25)
                emit_progress(ProcessingStep.EMBEDDING, embed_progress, f"向量化文字 {idx + 1}/{len(chunks)}...", chunk_count)

    # Process images using Jina CLIP API (75-90%)
    emit_progress(ProcessingStep.EMBEDDING, 75, "向量化圖片中...")
    for img_idx, img_data in enumerate(images_data):
        try:
            emb = embed_service.embed_image_from_base64(img_data["base64"])
            vector_store.add_image_chunk(
                document_id=document_id,
                document_name=filename,
                image_base64=img_data["base64"],
                embedding=emb,
                description=f"Image from page {img_data['page']}",
            )
            chunk_count += 1
            # Progress: 75-90% for image embedding
            if images_data:
                img_progress = 75 + int((img_idx + 1) / len(images_data) * 15)
                emit_progress(ProcessingStep.EMBEDDING, img_progress, f"向量化圖片 {img_idx + 1}/{len(images_data)}...", chunk_count)
        except Exception:
            continue

    # Step 4: Indexing (90-100%)
    emit_progress(ProcessingStep.INDEXING, 90, "建立索引中...")

    # Store document metadata
    vector_store.add_document(
        document_id=document_id,
        filename=filename,
        doc_type="pdf",
        file_size=len(file_content),
        chunk_count=chunk_count,
    )

    emit_progress(ProcessingStep.DONE, 100, f"處理完成！共 {chunk_count} 個區塊", chunk_count)

    return document_id, chunk_count


def process_word(
    file_content: bytes,
    filename: str,
    use_smart_chunking: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> tuple[str, int]:
    """
    Process a Word document (.docx) and store chunks in vector store.

    Args:
        file_content: Raw Word document bytes
        filename: Original filename
        use_smart_chunking: Whether to use smart chunking
        progress_callback: Optional callback for progress updates

    Returns: (document_id, chunk_count)
    """
    document_id = str(uuid.uuid4())
    chunk_count = 0

    def emit_progress(step: ProcessingStep, progress: int, message: str, chunks: Optional[int] = None):
        if progress_callback:
            try:
                progress_callback(step, progress, message, chunks)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")

    images_data = []

    # Save original file to storage
    try:
        file_storage.save_file(document_id, filename, file_content)
        logger.info(f"Saved original file for document {document_id}")
    except Exception as e:
        logger.warning(f"Failed to save original file: {e}")

    # Step 1: Parsing (0-25%)
    emit_progress(ProcessingStep.PARSING, 0, "開始解析 Word 文件...")

    # Extract text from Word document
    doc = DocxDocument(io.BytesIO(file_content))

    # Collect all text
    full_text = ""
    paragraphs = doc.paragraphs
    total_paras = len(paragraphs)
    for para_idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if text:
            full_text += text + "\n"
        # Progress: 0-15% for paragraph extraction
        if total_paras > 0:
            para_progress = int((para_idx + 1) / total_paras * 15)
            emit_progress(ProcessingStep.PARSING, para_progress, f"解析段落 {para_idx + 1}/{total_paras}...")

    # Extract tables (15-20%)
    emit_progress(ProcessingStep.PARSING, 15, "解析表格中...")
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                table_text.append(" | ".join(row_text))
        if table_text:
            full_text += "\n" + "\n".join(table_text) + "\n"

    # Extract images from Word document (20-25%)
    emit_progress(ProcessingStep.PARSING, 20, "提取圖片中...")
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_bytes = rel.target_part.blob
                image_base64 = base64.b64encode(image_bytes).decode("utf-8")
                images_data.append({
                    "base64": image_base64,
                })
            except Exception:
                continue
    emit_progress(ProcessingStep.PARSING, 25, f"解析完成，找到 {len(images_data)} 張圖片")

    # Step 2: Chunking (25-50%)
    emit_progress(ProcessingStep.CHUNKING, 25, "文件分塊中...")

    # Use smart chunking
    if use_smart_chunking and full_text.strip():
        chunks = smart_chunk(full_text, chunk_size=500, overlap=100)
    else:
        # Fallback
        chunks = [
            {"content": p.strip(), "metadata": {}}
            for p in full_text.split("\n\n")
            if p.strip() and len(p.strip()) > 30
        ]

    total_items = len(chunks) + len(images_data)
    emit_progress(ProcessingStep.CHUNKING, 50, f"分塊完成，共 {len(chunks)} 個文字區塊", len(chunks))

    # Step 3: Embedding (50-90%)
    emit_progress(ProcessingStep.EMBEDDING, 50, "向量化文字區塊中...")

    # Process text chunks with metadata
    if chunks:
        contents = [c["content"] for c in chunks]
        embeddings = embed_service.embed_texts(contents)
        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            vector_store.add_text_chunk(
                document_id=document_id,
                document_name=filename,
                content=chunk["content"],
                embedding=emb,
                chunk_index=idx,
                metadata=chunk.get("metadata", {}),
            )
            chunk_count += 1
            # Progress: 50-75% for text embedding
            if total_items > 0:
                embed_progress = 50 + int((idx + 1) / total_items * 25)
                emit_progress(ProcessingStep.EMBEDDING, embed_progress, f"向量化文字 {idx + 1}/{len(chunks)}...", chunk_count)

    # Process images using Jina CLIP API (75-90%)
    emit_progress(ProcessingStep.EMBEDDING, 75, "向量化圖片中...")
    for img_idx, img_data in enumerate(images_data):
        try:
            emb = embed_service.embed_image_from_base64(img_data["base64"])
            vector_store.add_image_chunk(
                document_id=document_id,
                document_name=filename,
                image_base64=img_data["base64"],
                embedding=emb,
                description=f"Image from {filename}",
            )
            chunk_count += 1
            # Progress: 75-90% for image embedding
            if images_data:
                img_progress = 75 + int((img_idx + 1) / len(images_data) * 15)
                emit_progress(ProcessingStep.EMBEDDING, img_progress, f"向量化圖片 {img_idx + 1}/{len(images_data)}...", chunk_count)
        except Exception:
            continue

    # Step 4: Indexing (90-100%)
    emit_progress(ProcessingStep.INDEXING, 90, "建立索引中...")

    # Store document metadata
    vector_store.add_document(
        document_id=document_id,
        filename=filename,
        doc_type="word",
        file_size=len(file_content),
        chunk_count=chunk_count,
    )

    emit_progress(ProcessingStep.DONE, 100, f"處理完成！共 {chunk_count} 個區塊", chunk_count)

    return document_id, chunk_count


def process_image(
    file_content: bytes,
    filename: str,
    progress_callback: Optional[ProgressCallback] = None,
) -> tuple[str, int]:
    """
    Process an image file and store in vector store.

    Args:
        file_content: Raw image bytes
        filename: Original filename
        progress_callback: Optional callback for progress updates

    Returns: (document_id, chunk_count)
    """
    document_id = str(uuid.uuid4())

    def emit_progress(step: ProcessingStep, progress: int, message: str, chunks: Optional[int] = None):
        if progress_callback:
            try:
                progress_callback(step, progress, message, chunks)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")

    # Save original file to storage
    try:
        file_storage.save_file(document_id, filename, file_content)
        logger.info(f"Saved original file for document {document_id}")
    except Exception as e:
        logger.warning(f"Failed to save original file: {e}")

    # Step 1: Parsing (0-25%)
    emit_progress(ProcessingStep.PARSING, 0, "開始處理圖片...")

    # Convert to base64
    image_base64 = base64.b64encode(file_content).decode("utf-8")
    emit_progress(ProcessingStep.PARSING, 25, "圖片解析完成")

    # Step 2: Chunking - minimal for images (25-50%)
    emit_progress(ProcessingStep.CHUNKING, 50, "圖片準備完成", 1)

    # Step 3: Embedding (50-90%)
    emit_progress(ProcessingStep.EMBEDDING, 50, "向量化圖片中...")

    # Embed image using Jina CLIP API
    emb = embed_service.embed_image_from_bytes(file_content)

    emit_progress(ProcessingStep.EMBEDDING, 75, "向量化完成")

    # Store in vector store
    vector_store.add_image_chunk(
        document_id=document_id,
        document_name=filename,
        image_base64=image_base64,
        embedding=emb,
        description=f"Uploaded image: {filename}",
    )

    emit_progress(ProcessingStep.EMBEDDING, 90, "圖片已儲存", 1)

    # Step 4: Indexing (90-100%)
    emit_progress(ProcessingStep.INDEXING, 90, "建立索引中...")

    # Store document metadata
    vector_store.add_document(
        document_id=document_id,
        filename=filename,
        doc_type="image",
        file_size=len(file_content),
        chunk_count=1,
    )

    emit_progress(ProcessingStep.DONE, 100, "處理完成！", 1)

    return document_id, 1


def process_excel(
    file_content: bytes,
    filename: str,
    use_smart_chunking: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> tuple[str, int]:
    """
    Process an Excel file (.xlsx or .xls) and store chunks in vector store.

    Each worksheet is converted to structured text preserving table format.
    Uses openpyxl for .xlsx and xlrd for .xls files.

    Args:
        file_content: Raw Excel file bytes
        filename: Original filename
        use_smart_chunking: Whether to use smart chunking
        progress_callback: Optional callback for progress updates

    Returns: (document_id, chunk_count)
    """
    document_id = str(uuid.uuid4())
    chunk_count = 0

    def emit_progress(step: ProcessingStep, progress: int, message: str, chunks: Optional[int] = None):
        if progress_callback:
            try:
                progress_callback(step, progress, message, chunks)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")

    # Save original file to storage
    try:
        file_storage.save_file(document_id, filename, file_content)
        logger.info(f"Saved original file for document {document_id}")
    except Exception as e:
        logger.warning(f"Failed to save original file: {e}")

    # Step 1: Parsing (0-25%)
    emit_progress(ProcessingStep.PARSING, 0, "開始解析 Excel 文件...")

    sheets_data = []
    is_xlsx = filename.lower().endswith(".xlsx")

    if is_xlsx:
        # Process .xlsx with openpyxl
        wb = load_workbook(io.BytesIO(file_content), data_only=True)
        sheet_names = wb.sheetnames
        total_sheets = len(sheet_names)

        for sheet_idx, sheet_name in enumerate(sheet_names):
            ws = wb[sheet_name]
            rows = []

            for row in ws.iter_rows(values_only=True):
                # Convert None values to empty strings
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):  # Skip completely empty rows
                    rows.append(row_values)

            if rows:
                sheets_data.append({
                    "name": sheet_name,
                    "rows": rows,
                })

            # Progress: 0-20% for parsing
            parse_progress = int((sheet_idx + 1) / total_sheets * 20)
            emit_progress(ProcessingStep.PARSING, parse_progress, f"解析工作表 {sheet_idx + 1}/{total_sheets}: {sheet_name}...")

        wb.close()
    else:
        # Process .xls with xlrd
        wb = xlrd.open_workbook(file_contents=file_content)
        total_sheets = wb.nsheets

        for sheet_idx in range(total_sheets):
            ws = wb.sheet_by_index(sheet_idx)
            sheet_name = ws.name
            rows = []

            for row_idx in range(ws.nrows):
                row_values = [str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols)]
                if any(v.strip() for v in row_values):  # Skip completely empty rows
                    rows.append(row_values)

            if rows:
                sheets_data.append({
                    "name": sheet_name,
                    "rows": rows,
                })

            # Progress: 0-20% for parsing
            parse_progress = int((sheet_idx + 1) / total_sheets * 20)
            emit_progress(ProcessingStep.PARSING, parse_progress, f"解析工作表 {sheet_idx + 1}/{total_sheets}: {sheet_name}...")

    emit_progress(ProcessingStep.PARSING, 25, f"解析完成，共 {len(sheets_data)} 個工作表")

    # Step 2: Chunking (25-50%)
    emit_progress(ProcessingStep.CHUNKING, 25, "文件分塊中...")

    chunks = []
    for sheet in sheets_data:
        sheet_name = sheet["name"]
        rows = sheet["rows"]

        if not rows:
            continue

        # Convert rows to structured text format (markdown table style)
        # First row is treated as header
        header = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        # Build structured text for the sheet
        sheet_text_parts = []
        sheet_text_parts.append(f"## 工作表: {sheet_name}\n")

        if header:
            # Create markdown table header
            header_line = "| " + " | ".join(header) + " |"
            separator = "| " + " | ".join(["---"] * len(header)) + " |"
            sheet_text_parts.append(header_line)
            sheet_text_parts.append(separator)

        # Add data rows
        for row in data_rows:
            # Ensure row has same number of columns as header
            padded_row = row + [""] * (len(header) - len(row)) if len(row) < len(header) else row[:len(header)]
            row_line = "| " + " | ".join(padded_row) + " |"
            sheet_text_parts.append(row_line)

        sheet_text = "\n".join(sheet_text_parts)

        # Smart chunking or fallback
        if use_smart_chunking and sheet_text.strip():
            sheet_chunks = smart_chunk(sheet_text, chunk_size=500, overlap=100)
            for chunk in sheet_chunks:
                chunk["metadata"]["sheet_name"] = sheet_name
                chunks.append(chunk)
        else:
            # Fallback: treat entire sheet as one chunk
            if sheet_text.strip():
                chunks.append({
                    "content": sheet_text.strip(),
                    "metadata": {"sheet_name": sheet_name},
                })

    emit_progress(ProcessingStep.CHUNKING, 50, f"分塊完成，共 {len(chunks)} 個文字區塊", len(chunks))

    # Step 3: Embedding (50-90%)
    emit_progress(ProcessingStep.EMBEDDING, 50, "向量化文字區塊中...")

    if chunks:
        contents = [c["content"] for c in chunks]
        embeddings = embed_service.embed_texts(contents)

        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            vector_store.add_text_chunk(
                document_id=document_id,
                document_name=filename,
                content=chunk["content"],
                embedding=emb,
                chunk_index=idx,
                metadata=chunk.get("metadata", {}),
            )
            chunk_count += 1

            # Progress: 50-90% for embedding
            if len(chunks) > 0:
                embed_progress = 50 + int((idx + 1) / len(chunks) * 40)
                emit_progress(ProcessingStep.EMBEDDING, embed_progress, f"向量化文字 {idx + 1}/{len(chunks)}...", chunk_count)

    # Step 4: Indexing (90-100%)
    emit_progress(ProcessingStep.INDEXING, 90, "建立索引中...")

    # Store document metadata
    vector_store.add_document(
        document_id=document_id,
        filename=filename,
        doc_type="excel",
        file_size=len(file_content),
        chunk_count=chunk_count,
    )

    emit_progress(ProcessingStep.DONE, 100, f"處理完成！共 {chunk_count} 個區塊", chunk_count)

    return document_id, chunk_count


def get_file_type(filename: str) -> Optional[str]:
    """Get file type from filename."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    elif lower.endswith((".docx", ".doc")):
        return "word"
    elif lower.endswith((".xlsx", ".xls")):
        return "excel"
    elif lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
        return "image"
    return None
